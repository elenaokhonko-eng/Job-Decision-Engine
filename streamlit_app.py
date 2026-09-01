import os
import sys
import subprocess
import urllib.request
import json
import psycopg2
from psycopg2.extras import RealDictCursor
import datetime
import pandas as pd
import streamlit as st
import docx
from fpdf import FPDF
from io import BytesIO

# Load environment variables from Streamlit secrets, .env, and .env.local
def load_dotenv():
    # Load from Streamlit Cloud Secrets if available
    try:
        if hasattr(st, "secrets") and st.secrets:
            for key, val in st.secrets.items():
                if isinstance(val, str):
                    os.environ[key] = val
    except Exception:
        pass

    # Load from local .env and .env.local files
    script_dir = os.path.dirname(os.path.abspath(__file__))
    for name in [".env", ".env.local"]:
        filename = os.path.join(script_dir, name)
        if os.path.exists(filename):
            with open(filename, "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if line and not line.startswith("#") and "=" in line:
                        key, val = line.split("=", 1)
                        val = val.strip("'\"")
                        os.environ[key.strip()] = val

load_dotenv()


# Configure the page setting with modern style
st.set_page_config(
    page_title="Job Decision Engine - High-Autonomy Career Architect",
    page_icon="💼",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom dark-theme styling for professional visual aesthetics
st.markdown("""
<style>
    .reportview-container {
        background-color: #0F0F0F;
    }
    .metric-card {
        background-color: #161616;
        border: 1px solid #2A2A2A;
        padding: 15px;
        border-radius: 10px;
        text-align: center;
    }
    .title-accent {
        color: #D4AF37;
        font-family: 'serif';
    }
    .disclaimer {
        font-family: monospace;
        font-size: 11px;
        color: #888888;
    }
    .top-rec-card {
        background-color: #1a1a24;
        color: #e0e0e0;
        border-left: 5px solid #22c55e;
        padding: 15px;
        margin-bottom: 12px;
        border-radius: 6px;
    }
    .top-rec-card h4 {
        color: #ffffff;
        margin-top: 0;
        margin-bottom: 8px;
    }
    .top-rec-card p {
        color: #cccccc;
        margin-bottom: 4px;
        font-size: 14px;
    }
</style>
""", unsafe_allow_html=True)

# Neon Database Helper
def get_db_connection():
    database_url = os.environ.get("DATABASE_URL")
    if not database_url:
        st.error("❌ DATABASE_URL environment variable is missing. Please set it in your environment or Streamlit Secrets.")
        st.stop()
    
    # Ensure sslmode=require for Neon serverless Postgres
    if "sslmode=" not in database_url and "localhost" not in database_url and "127.0.0.1" not in database_url:
        if "?" in database_url:
            database_url += "&sslmode=require"
        else:
            database_url += "?sslmode=require"
            
    return psycopg2.connect(database_url)

def run_checked_command(command_args, step_label):
    """Run a local command safely and surface stdout/stderr to the UI."""
    result = subprocess.run(
        command_args,
        capture_output=True,
        text=True,
        check=True,
        shell=False,
        env=os.environ,
    )
    if result.stdout:
        st.code(result.stdout, language="text")
    if result.stderr:
        st.warning(f"{step_label} stderr:\n{result.stderr}")
    return result

def fetch_jobs_from_db():
    """
    Fetch the canonical shortlist from v_canonical_shortlist (migration 007).
    """
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        cursor.execute("""
            SELECT
                canonical_job_id        AS id,
                job_version_id,
                title,
                company,
                canonical_url           AS careers_portal_url,
                location,
                workplace_type,
                employment_type,
                description,
                gate_status,
                rejection_codes,
                gate_evidence_quotes,
                primary_lane            AS assigned_track,
                secondary_lanes,
                lane_confidence,
                priority_score          AS confidence_level,
                processing_status       AS status,
                nd_friendly_score,
                politics_stress_score,
                sensory_overload_index,
                next_action,
                strategic_value,
                recommended_cv_version,
                evaluation_summary,
                eval_provider,
                eval_is_fallback,
                version_mismatch,
                observed_at::text       AS "postedDate",
                evaluated_at,
                lane_matches,
                workability_facts,
                queue_status
            FROM v_canonical_shortlist
            ORDER BY observed_at DESC
        """)
        rows = cursor.fetchall()
        cursor.close()
        conn.close()
        return [dict(r) for r in rows]
    except Exception as e:
        st.error(f"Failed to fetch jobs from database: {e}")
        return []

def fetch_rejected_jobs_from_db():
    """Fetch hard-rejected and removed jobs for audit inspection."""
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        cursor.execute("""
            SELECT
                c.id,
                c.normalized_title AS title,
                c.company_name AS company,
                c.canonical_url AS careers_portal_url,
                c.processing_status AS status,
                c.rejection_reason,
                gd.decision AS gate_status,
                gd.rejection_codes,
                gd.evidence_quotes,
                c.created_at::text AS "postedDate"
            FROM canonical_jobs c
            LEFT JOIN (
                SELECT DISTINCT ON (canonical_job_id) canonical_job_id, decision, rejection_codes, evidence_quotes
                FROM gate_decisions ORDER BY canonical_job_id, created_at DESC
            ) gd ON gd.canonical_job_id = c.id
            WHERE c.processing_status IN ('HARD_REJECTED', 'MANUALLY_REMOVED')
            ORDER BY c.created_at DESC
            LIMIT 50
        """)
        rows = cursor.fetchall()
        cursor.close()
        conn.close()
        return [dict(r) for r in rows]
    except Exception as e:
        return []

def delete_job_from_db(job_id):
    """Soft-delete a canonical job by marking it MANUALLY_REMOVED.
    HARD_REJECTED is reserved for deterministic gate failures only (AGENTS.md invariant).
    """
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Soft-delete: set status to MANUALLY_REMOVED (never hard-delete; preserves audit trail)
        cursor.execute(
            """
            UPDATE canonical_jobs
            SET processing_status = 'MANUALLY_REMOVED',
                rejection_reason = 'Manually removed by user via Streamlit UI',
                updated_at = NOW()
            WHERE id = %s
            """,
            (job_id,)
        )

        conn.commit()
        cursor.close()
        conn.close()
        st.success("Listing removed successfully (soft-delete — record preserved for audit).")
        return True
    except Exception as e:
        st.error(f"Failed to remove job: {e}")
        return False

def save_new_job_to_db(job):
    """Insert a manually-added job into raw_job_observations so it enters the
    canonical pipeline (normalize → gate → lane → budget → evaluate).
    Writing to raw_jobs (legacy) is insufficient — that table is not read by the pipeline.
    """
    import hashlib, uuid
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        description_text = job.get("description", "")
        payload = {
            "company_name": job.get("company", "Unknown"),
            "title": job.get("title", "Unknown"),
            "description": description_text,
            "source": job.get("source", "MANUAL_STREAMLIT"),
            "careers_portal_url": job.get("careers_portal_url", ""),
        }
        raw_payload_str = json.dumps(payload)
        raw_payload_hash = hashlib.sha256(raw_payload_str.encode()).hexdigest()

        # Ensure a source_run exists for manual entries
        cursor.execute(
            "INSERT INTO source_runs (status) VALUES ('MANUAL_STREAMLIT') RETURNING id"
        )
        source_run_id = cursor.fetchone()[0]

        cursor.execute(
            """
            INSERT INTO raw_job_observations (
                source_run_id, source_name, source_external_id, source_url,
                retrieved_at, company_name, title, description_raw,
                location_raw, workplace_type_raw, employment_type_raw, compensation_raw,
                canonical_apply_url, source_lane, search_plan_version,
                raw_payload, raw_payload_hash
            ) VALUES (%s, %s, %s, %s, NOW(), %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            ON CONFLICT (raw_payload_hash) DO NOTHING
            """,
            (
                source_run_id,
                "MANUAL_STREAMLIT",
                f"manual-{raw_payload_hash[:16]}",
                job.get("careers_portal_url", ""),
                job.get("company", "Unknown"),
                job.get("title", "Unknown"),
                description_text,
                job.get("location", "Singapore"),
                "UNKNOWN",
                "UNKNOWN",
                job.get("salaryRange", "UNKNOWN"),
                job.get("careers_portal_url", ""),
                "UNKNOWN",
                "1.0",
                raw_payload_str,
                raw_payload_hash,
            ),
        )

        conn.commit()
        cursor.close()
        conn.close()
        st.success(f"✅ '{job['title']}' added to the ingestion pipeline (raw_job_observations). It will be normalized, gated, and evaluated in the next run.")
        return True
    except Exception as e:
        st.error(f"Failed to save job: {e}")
        return False

def convert_markdown_to_docx(md_text):
    doc = docx.Document()
    for line in md_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("* ") or line.startswith("- "):
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:]
            parts = text.split("**")
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                if idx % 2 == 1:
                    run.bold = True
        else:
            p = doc.add_paragraph()
            parts = line.split("**")
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                if idx % 2 == 1:
                    run.bold = True
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

class PDFResume(FPDF):
    def header(self):
        pass
    def footer(self):
        # Removed page numbers for ATS compatibility
        pass

def convert_markdown_to_pdf(md_text):
    # Sanitize Unicode characters that helvetica (latin-1) cannot encode
    replacements = {
        '\u2013': '-', '\u2014': '-', '\u2018': "'", '\u2019': "'", 
        '\u201c': '"', '\u201d': '"', '\u2022': '*', '\u00A0': ' ',
        '\u2026': '...'
    }
    for k, v in replacements.items():
        md_text = md_text.replace(k, v)
    # Strip any remaining non-latin1 characters
    md_text = md_text.encode('latin-1', 'ignore').decode('latin-1')
    
    import re
    # Break extremely long words/URLs that crash FPDF multi_cell (40 chars to be ultra safe)
    md_text = re.sub(r'(\S{40})', r'\1 ', md_text)
    
    pdf = PDFResume()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("helvetica", size=10)
    
    for line in md_text.split("\n"):
        line = line.strip()
        if not line:
            pdf.ln(4)
            continue
            
        try:
            # Force cursor to left margin to prevent 'Not enough horizontal space' errors
            pdf.set_x(pdf.l_margin)
            if line.startswith("# "):
                pdf.set_font("helvetica", "B", 16)
                pdf.multi_cell(0, 10, line[2:])
                pdf.ln(2)
            elif line.startswith("## "):
                pdf.set_font("helvetica", "B", 13)
                pdf.multi_cell(0, 8, line[3:])
                pdf.ln(1)
            elif line.startswith("### "):
                pdf.set_font("helvetica", "B", 11)
                pdf.multi_cell(0, 6, line[4:])
                pdf.ln(1)
            elif line.startswith("* ") or line.startswith("- "):
                pdf.set_font("helvetica", "", 10)
                text = line[2:]
                text_clean = text.replace("**", "")
                pdf.multi_cell(0, 5, f"-  {text_clean}")
            else:
                pdf.set_font("helvetica", "", 10)
                text_clean = line.replace("**", "")
                pdf.multi_cell(0, 5, text_clean)
        except Exception as e:
            # Fallback for lines that completely break FPDF
            pdf.set_font("helvetica", "", 10)
            pdf.multi_cell(0, 5, "[Error rendering line]")
            
    return bytes(pdf.output())

def python_generate_content(contents, system_instruction=None, response_mime_type=None, response_schema=None):
    # Load keys
    gemini_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GEMINI_FLASH_API_KEY")
    openai_key = os.environ.get("OPENAI_API_KEY")
    
    # 1. Try OpenAI models in fallback sequence
    if openai_key:
        models_to_try = [
            os.environ.get("OPENAI_MODEL", "gpt-5.6-sol"),
            "gpt-5.6-terra",
            "o3-mini",
            "gpt-4o"
        ]
        
        # Remove duplicates preserving order
        unique_models = []
        for m in models_to_try:
            if m not in unique_models:
                unique_models.append(m)
                
        openai_success = False
        openai_text = ""
        
        for model_name in unique_models:
            try:
                url = "https://api.openai.com/v1/chat/completions"
                headers = {
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {openai_key}"
                }
                messages = []
                if system_instruction:
                    messages.append({"role": "system", "content": system_instruction})
                messages.append({"role": "user", "content": contents})
                
                body = {
                    "model": model_name,
                    "messages": messages,
                    "temperature": 0.0,
                    "max_completion_tokens": 16384
                }
                if response_schema:
                    body["response_format"] = {
                        "type": "json_schema",
                        "json_schema": {
                            "name": "structured_response",
                            "schema": response_schema,
                            "strict": True
                        }
                    }
                elif response_mime_type == "application/json":
                    body["response_format"] = {"type": "json_object"}
                    
                req = urllib.request.Request(url, data=json.dumps(body).encode("utf-8"), headers=headers, method="POST")
                with urllib.request.urlopen(req, timeout=90) as response:
                    res_data = json.loads(response.read().decode("utf-8"))
                    openai_text = res_data["choices"][0]["message"]["content"]
                    openai_success = True
                    break # Break out of the loop on success
            except Exception as openai_err:
                st.warning(f"⚠️ OpenAI model {model_name} failed: {openai_err}. Trying next fallback...")
                continue
                
        if openai_success:
            return openai_text

    # 2. Try Gemini second
    if gemini_key:
        try:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-3.6-flash:generateContent?key={gemini_key}"
            headers = {"Content-Type": "application/json"}
            
            body = {
                "contents": [{"parts": [{"text": contents}]}],
                "generationConfig": {}
            }
            if response_schema:
                body["generationConfig"]["responseSchema"] = response_schema
                body["generationConfig"]["responseMimeType"] = "application/json"
            elif response_mime_type:
                body["generationConfig"]["responseMimeType"] = response_mime_type
                
            if system_instruction:
                body["systemInstruction"] = {"parts": [{"text": system_instruction}]}
                
            req = urllib.request.Request(url, data=json.dumps(body).encode("utf-8"), headers=headers, method="POST")
            with urllib.request.urlopen(req, timeout=90) as response:
                res_data = json.loads(response.read().decode("utf-8"))
                text = res_data["candidates"][0]["content"]["parts"][0]["text"]
                return text
        except Exception as gemini_err:
            st.warning(f"⚠️ Gemini request failed: {gemini_err}.")
            
    raise Exception("All configured models (OpenAI, Gemini) failed or no API keys are set.")

def ingest_linkedin_saved_json(jobs):
    import hashlib
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Ensure a source_run exists
        cursor.execute("INSERT INTO source_runs (status) VALUES ('LINKEDIN_IMPORT') RETURNING id")
        source_run_id = cursor.fetchone()[0]
        
        inserted_count = 0
        skipped_count = 0
        
        for job in jobs:
            title = job.get("title", "").strip()
            company = job.get("company", "").strip()
            url = job.get("url", "").strip()
            description = job.get("description", "").strip()
            location = job.get("location", "Singapore").strip()
            
            if not title or not company or not url or not description:
                skipped_count += 1
                continue
                
            raw_payload_str = json.dumps(job)
            raw_payload_hash = hashlib.sha256(raw_payload_str.encode()).hexdigest()
            
            cursor.execute("""
                INSERT INTO raw_job_observations (
                    source_run_id, source_name, source_external_id, source_url,
                    retrieved_at, company_name, title, description_raw,
                    location_raw, workplace_type_raw, employment_type_raw, compensation_raw,
                    canonical_apply_url, source_lane, search_plan_version,
                    raw_payload, raw_payload_hash
                ) VALUES (%s, 'LINKEDIN', %s, %s, NOW(), %s, %s, %s, %s, 'UNKNOWN', 'PERMANENT', 'UNKNOWN', %s, 'UNKNOWN', '1.0', %s, %s)
                ON CONFLICT (raw_payload_hash) DO NOTHING
                RETURNING id
            """, (
                source_run_id,
                f"linkedin-{raw_payload_hash[:16]}",
                url,
                company,
                title,
                description,
                location,
                url,
                raw_payload_str,
                raw_payload_hash
            ))
            res = cursor.fetchone()
            if res:
                inserted_count += 1
            else:
                skipped_count += 1
                
        conn.commit()
        cursor.close()
        conn.close()
        return inserted_count, skipped_count
    except Exception as e:
        st.error(f"Database insertion failed: {e}")
        return 0, 0

def fetch_company_analytics_from_db():
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        cursor.execute("""
            SELECT name as "Company", 
                   industry as "Industry",
                   nd_friendly_avg_score as "Avg Autonomy Score",
                   politics_stress_avg_score as "Avg Politics Score",
                   sensory_overload_avg_index as "Avg Sensory Index",
                   focus_protection_avg_score as "Avg Focus Score",
                   is_neurodivergent_approved as "Approved",
                   is_toxic_culture_blacklisted as "Toxic"
            FROM companies
            ORDER BY nd_friendly_avg_score DESC
        """)
        rows = cursor.fetchall()
        cursor.close()
        conn.close()
        return [dict(r) for r in rows]
    except Exception as e:
        st.error(f"Failed to fetch company analytics: {e}")
        return []

# Fetch data
jobs_list = fetch_jobs_from_db()

# Title
st.title("💼 Job Decision Engine — Streamlit Console (v4.1)")
st.markdown("### *Multi-Stage Weighted High-Autonomy Technical Architect & Builder Console*")
st.markdown("---")

# Sidebar - Filters & Stats
st.sidebar.header("🎯 Navigation & Filters")

# Metrics
total_jobs = len(jobs_list)
evaluated_count = sum(1 for j in jobs_list if j.get("status") and j.get("status") != "UNASSIGNED")
approved_count = sum(1 for j in jobs_list if j.get("status") in ("STRONG MATCH", "PRIORITY_APPLY", "HIGH_FIT_HIGH_RISK"))
review_count = sum(1 for j in jobs_list if j.get("status") in ("REVIEW REQUIRED", "APPLY_AFTER_VERIFICATION"))
toxic_count = sum(1 for j in jobs_list if (j.get("politics_stress_score") or 0) >= 70 or (j.get("nd_friendly_score") or 100) < 50)

st.sidebar.subheader("📊 Engine Statistics")
st.sidebar.metric("Total Vault Jobs", total_jobs)
st.sidebar.metric("Fully Evaluated", evaluated_count)
st.sidebar.metric("Top Recommended (Strong)", approved_count)
st.sidebar.metric("Needs Review", review_count)
st.sidebar.metric("Toxicity Flags", toxic_count)

st.sidebar.markdown("---")
st.sidebar.subheader("📅 Automated Schedules")
st.sidebar.info("""
* **Daily Ingestion & Evaluation**: Runs daily at **10:00 AM SGT** (02:00 UTC) via GitHub Actions.
* **Weekly LinkedIn Auto-Sync**: Runs every **Sunday at 10:00 AM SGT** (02:00 UTC) via GitHub Actions.
""")

st.sidebar.subheader("⚡ Pipeline Action Controls")

is_local = os.path.exists(".env.local")
github_token = os.environ.get("GITHUB_TOKEN") or os.environ.get("GH_TOKEN") or os.environ.get("GITHUB_PAT")

# Button 1: Ingest & Process Pipeline (GitHub Actions or Local)
if st.sidebar.button("⚡ Run Job Discovery & Evaluation Pipeline", help="Runs the full pipeline: Ingest Gmail & ATS adapters, Normalize, Hard Gate, Semantic Lane Route, Budget, and AI Evaluate."):
    if not is_local and not github_token:
        st.sidebar.error("⚠️ GITHUB_TOKEN is missing in Streamlit secrets. Please configure it to trigger GitHub Action workflows from the cloud.")
    else:
        if github_token:
            with st.spinner("Triggering GitHub Actions ingest.yml workflow..."):
                try:
                    req = urllib.request.Request(
                        "https://api.github.com/repos/elenaokhonko-eng/Job-Decision-Engine/actions/workflows/ingest.yml/dispatches",
                        data=json.dumps({"ref": "main"}).encode("utf-8"),
                        headers={"Authorization": f"Bearer {github_token}", "Accept": "application/vnd.github.v3+json", "User-Agent": "StreamlitConsole"},
                        method="POST"
                    )
                    with urllib.request.urlopen(req) as resp:
                        if resp.status in (204, 200, 201):
                            st.success("🐙 Triggered GitHub Actions Job Discovery Ingestion (ingest.yml) workflow!")
                            st.balloons()
                except Exception as gh_err:
                    st.error(f"GitHub Trigger Error: {gh_err}")
        else:
            with st.spinner("Running full pipeline locally..."):
                try:
                    st.info("Step 1/4: Ingesting Gmail alerts...")
                    run_checked_command(["npx", "tsx", "scripts/ingest_gmail.ts"], "ingest_gmail")
                    
                    st.info("Step 2/4: Polling ATS & Job Board adapters...")
                    run_checked_command(["npx", "tsx", "scripts/run_adapters.ts"], "run_adapters")
                    
                    st.info("Step 3/4: Parsing email alerts & staging observations...")
                    run_checked_command(["npx", "tsx", "scripts/parse_emails.ts"], "parse_emails")

                    st.info("Step 4/4: Running Discovery Pipeline (Normalize, Gate, Route, Budget)...")
                    run_checked_command(["npx", "tsx", "scripts/process_pipeline.ts"], "process_pipeline")

                    st.info("Step 5/5: Running AI Evaluation Queue Processor...")
                    run_checked_command(["npx", "tsx", "scripts/evaluate_queue.ts"], "evaluate_queue")
                    st.success("✅ Full pipeline execution finished!")
                    st.balloons()
                except subprocess.CalledProcessError as cpe:
                    st.error(f"Execution failed at step with exit code {cpe.returncode}.")
                    if cpe.stdout:
                        st.code(cpe.stdout, language="text")
                    if cpe.stderr:
                        st.error(cpe.stderr)
                except Exception as e:
                    st.error(f"Execution Error: {e}")
        st.rerun()

st.sidebar.markdown("---")
st.sidebar.subheader("🔍 Filter Listings")
search_query = st.sidebar.text_input("Keyword Search", "")
lane_filter = st.sidebar.selectbox("Filter Target Lane", ["All Lanes", "CORE_AI_DATA", "LEGAL_REGTECH", "HEALTH_BIO_PHARMA", "INVESTMENT_MARKETS_FINTECH", "UNCLASSIFIED"])
status_filter = st.sidebar.selectbox("Filter Pipeline Status", ["All Statuses", "AI_EVALUATED", "QUEUED_FOR_AI", "LANE_ROUTED", "PREQUALIFIED", "NEEDS_VERIFICATION", "DEFERRED_BUDGET", "ROUTING_DEFERRED"])
source_values = sorted({(j.get("source") or "UNKNOWN") for j in jobs_list})
board_filter = st.sidebar.selectbox("Filter Source", ["All Sources", *source_values])
track_values = sorted({(j.get("assigned_track") or "UNASSIGNED") for j in jobs_list})
track_filter = st.sidebar.selectbox("Filter Track", ["All Tracks", "Unassigned", *track_values])

# Apply filters
filtered_jobs = jobs_list
if lane_filter and lane_filter != "All Lanes":
    filtered_jobs = [j for j in filtered_jobs if j.get("assigned_track") == lane_filter]
if status_filter and status_filter != "All Statuses":
    filtered_jobs = [j for j in filtered_jobs if j.get("status") == status_filter]
if search_query:
    filtered_jobs = [j for j in filtered_jobs if search_query.lower() in (j.get("title") or "").lower() or search_query.lower() in (j.get("company") or "").lower() or search_query.lower() in (j.get("description") or "").lower()]
if board_filter != "All Sources":
    filtered_jobs = [j for j in filtered_jobs if (j.get("source") or "UNKNOWN") == board_filter]
if track_filter != "All Tracks":
    if track_filter == "Unassigned":
        filtered_jobs = [j for j in filtered_jobs if not j.get("assigned_track") or j.get("assigned_track") == "UNASSIGNED"]
    else:
        filtered_jobs = [j for j in filtered_jobs if j.get("assigned_track") == track_filter]

# Main Dashboard Layout tabs
tab_dashboard, tab_add_job, tab_linkedin, tab_analytics, tab_cv = st.tabs(["📁 Postgres Job Vault", "➕ Add Job Ad", "🔗 LinkedIn Saved Jobs", "🔥 ND Culture Analytics", "📄 CV Customizer"])

with tab_dashboard:
    # Segment out Top Recommended Jobs based on next_action and AI evaluation
    top_recommended = [j for j in jobs_list if j.get("next_action") in ("PRIORITY_APPLY", "APPLY_NOW", "CUSTOMIZE_CV", "APPLY_AFTER_VERIFICATION") or j.get("status") == "AI_EVALUATED"]
    top_recommended = sorted(top_recommended, key=lambda x: (x.get("nd_friendly_score") or 0), reverse=True)[:10]

    st.subheader("🏆 Top Recommended Opportunities")
    if not top_recommended:
        st.info("No priority evaluated recommendations found yet. Run the discovery & evaluation pipeline.")
    else:
        cols = st.columns(2)
        for idx, rjob in enumerate(top_recommended):
            col_idx = idx % 2
            with cols[col_idx]:
                version_warn = " ⚠️ Stale Evaluation" if rjob.get("version_mismatch") else ""
                st.markdown(f"""
                <div class="top-rec-card">
                    <h4>⭐ #{idx+1} {rjob['title']} {version_warn}</h4>
                    <p><b>Company:</b> {rjob['company']} | <b>Lane:</b> <code>{rjob.get('assigned_track')}</code></p>
                    <p><b>Location:</b> {rjob.get('location')} ({rjob.get('workplace_type')})</p>
                    <p><b>Action:</b> <code style='color:#22c55e;'>{rjob.get('next_action')}</code> | <b>Confidence:</b> {rjob.get('lane_confidence')}</p>
                    <p><b>Workplace Culture:</b> Autonomy: {rjob.get('nd_friendly_score')}% | Politics: {rjob.get('politics_stress_score')}%</p>
                    <p><b>Summary:</b> {rjob.get('evaluation_summary') or 'N/A'}</p>
                </div>
                """, unsafe_allow_html=True)
                st.markdown(f"🔗 [Verify Job Ad & Apply]({rjob['careers_portal_url']})")

    st.markdown("---")

    # Interactive search filters inside the main tab
    search_col1, search_col2 = st.columns(2)
    with search_col1:
        search_title = st.text_input("🔍 Search Job Title", "", key="search_title_main")
    with search_col2:
        search_company = st.text_input("🏢 Search Company Name", "", key="search_company_main")

    if search_title:
        filtered_jobs = [j for j in filtered_jobs if search_title.lower() in (j.get("title") or "").lower()]
    if search_company:
        filtered_jobs = [j for j in filtered_jobs if search_company.lower() in (j.get("company") or "").lower()]

    col_left, col_right = st.columns([2, 3])

    with col_left:
        # Sort and split active vs rejected jobs
        def status_sort_key(j):
            status = j.get("status", "UNASSIGNED")
            score = (j.get("total_score") or 0)
            if status in ("STRONG MATCH", "PRIORITY_APPLY", "HIGH_FIT_HIGH_RISK"):
                return (0, -score)
            elif status in ("REVIEW REQUIRED", "APPLY_AFTER_VERIFICATION"):
                return (1, -score)
            elif status == "REJECTED":
                return (2, -score)
            else:
                return (3, -score)

        sorted_filtered_jobs = sorted(filtered_jobs, key=status_sort_key)
        active_jobs = [j for j in sorted_filtered_jobs if j.get("status") != "REJECTED"]
        rejected_jobs = [j for j in sorted_filtered_jobs if j.get("status") == "REJECTED"]

        st.subheader("📋 Available Listings Vault")
        if not active_jobs and not rejected_jobs:
            st.info("No matching jobs in the current Postgres database.")
        else:
            # Render Active (Green & Orange) Jobs
            if active_jobs:
                st.write(f"Showing {len(active_jobs)} Active Match Listings:")
                for idx, job in enumerate(active_jobs):
                    status = job.get("status", "UNASSIGNED")
                    score = job.get("total_score", 0)
                    company = job.get("company", "Unknown")
                    title = job.get("title", "Job Title")
                    badge_style = "✅" if status in ("STRONG MATCH", "PRIORITY_APPLY", "HIGH_FIT_HIGH_RISK") else "⚠️"
                    
                    with st.expander(f"{badge_style} {title} — {company} ({status})"):
                        st.markdown(f"**Source Board:** `{job.get('source')}`")
                        st.markdown(f"**Salary Range:** {job.get('salaryRange') or 'Not specified'}")
                        st.markdown(f"**Location:** {job.get('location') or 'Singapore'}")
                        st.markdown(f"**Verification Link:** [Go to Careers Portal]({job.get('careers_portal_url')})")
                        score = job.get('total_score', 0)
                        aut = job.get('nd_friendly_score')
                        pol = job.get('politics_stress_score')
                        
                        aut_str = f"🟢 {aut}%" if aut and aut >= 70 else (f"🔴 {aut}%" if aut else "N/A")
                        pol_str = f"🔴 {pol}%" if pol and pol >= 40 else (f"🟢 {pol}%" if pol else "N/A")
                        
                        st.markdown(f"**Match Score:** `{score}/100`")
                        st.markdown(f"**Autonomy Score:** {aut_str} | **Politics Stress:** {pol_str}")
                        
                        desc_text = job.get("description", "")
                        parsed_desc = None
                        if isinstance(desc_text, dict):
                            parsed_desc = desc_text
                            desc_text = parsed_desc.get("job_description", "")
                        elif isinstance(desc_text, str) and desc_text.strip().startswith("{"):
                            try:
                                parsed_desc = json.loads(desc_text)
                                desc_text = parsed_desc.get("job_description", "")
                            except Exception:
                                pass
                        st.text_area("Full Description Brief", desc_text or "", height=100, disabled=True, key=f"active_desc_{idx}")
                        
                        if st.button("🗑️ Delete Listing", key=f"active_del_{job.get('id') or idx}"):
                            if delete_job_from_db(job.get("id")):
                                st.rerun()
            else:
                if rejected_jobs:
                    st.info(f"💡 No active matches found, but {len(rejected_jobs)} matching listings are in the Rejected/Discarded folder below.")
                else:
                    st.info("No active matching jobs (STRONG MATCH or REVIEW REQUIRED) are currently stored.")

            # Render Rejected (Red) Jobs inside an expander
            if rejected_jobs:
                st.markdown("---")
                is_search_active = bool(search_title or search_company or search_query)
                with st.expander(f"🔴 View Rejected/Discarded Listings ({len(rejected_jobs)} jobs)", expanded=is_search_active):
                    # Limit rendering of rejected popovers to top 50 to prevent severe browser lag
                    display_limit = 50
                    for idx, job in enumerate(rejected_jobs[:display_limit]):
                        company = job.get("company", "Unknown")
                        title = job.get("title", "Job Title")
                        score = (job.get("total_score") or 0)
                        
                        with st.popover(f"🔴 {title} — {company}"):
                            st.markdown(f"**Source Board:** `{job.get('source')}`")
                            st.markdown(f"**Verification Link:** [Go to Careers Portal]({job.get('careers_portal_url')})")
                            st.markdown(f"**Rejected Score:** `{score}/100`")
                            
                            desc_text = job.get("description", "")
                            parsed_desc = None
                            if isinstance(desc_text, dict):
                                parsed_desc = desc_text
                                desc_text = parsed_desc.get("job_description", "")
                            elif isinstance(desc_text, str) and desc_text.strip().startswith("{"):
                                try:
                                    parsed_desc = json.loads(desc_text)
                                    desc_text = parsed_desc.get("job_description", "")
                                except Exception:
                                    pass
                            st.text_area("Full Description Brief", desc_text or "", height=100, disabled=True, key=f"rej_desc_{idx}")
                            
                            if st.button("🗑️ Delete Listing Permanent", key=f"rej_del_{job.get('id') or idx}"):
                                try:
                                    conn = get_db_connection()
                                    cursor = conn.cursor()
                                    cursor.execute("DELETE FROM jobs WHERE id = %s", (job.get('id'),))
                                    conn.commit()
                                    cursor.close()
                                    conn.close()
                                    st.success("Permanently deleted!")
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"Failed to permanently delete: {e}")
                    if len(rejected_jobs) > display_limit:
                        st.caption(f"⚠️ Showing first {display_limit} rejected listings to maintain UI performance. Use the search inputs above to filter down further.")

    with col_right:
        st.subheader("🤖 Scoring & Match Analysis Details")
        st.write("Select an evaluated job to view detailed autonomy & focus match metrics, workplace stress assessments, and strategic CV targeting.")
        
        evaluated_jobs = [j for j in filtered_jobs if j.get("status") and j.get("status") not in ("UNASSIGNED", "REJECTED")]
        
        def format_job_option(j):
            return f"{j.get('title')} ({j.get('company')}) - {str(j.get('id'))[:8]}"
            
        selected_job_title = st.selectbox(
            "Select Job to Analyze", 
            [format_job_option(j) for j in evaluated_jobs] if evaluated_jobs else ["No Evaluated Jobs Available"]
        )
        
        # Get actual job object
        job_to_show = None
        if evaluated_jobs and selected_job_title != "No Evaluated Jobs Available":
            idx_selected = [format_job_option(j) for j in evaluated_jobs].index(selected_job_title)
            job_to_show = evaluated_jobs[idx_selected]

        if job_to_show:
            st.markdown(f"#### Selected: **{job_to_show['title']}** at *{job_to_show['company']}*")
            st.markdown(f"**Verifiable Careers Link:** `{job_to_show['careers_portal_url']}`")
            
            # Show score metrics
            st.markdown("---")
            core = job_to_show.get('total_score', 0)
            core_emoji = "🟢" if core >= 80 else "🔴"
            st.markdown(f"### Match Score: `{core} / 100` {core_emoji}")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                aut = job_to_show.get('nd_friendly_score')
                aut_emoji = "🟢" if aut and aut >= 70 else ("🔴" if aut else "")
                st.metric("Autonomy Culture Score", f"{aut}% {aut_emoji}")
            with col2:
                pol = job_to_show.get('politics_stress_score')
                pol_emoji = "🔴" if pol and pol >= 40 else ("🟢" if pol else "")
                st.metric("Politics Stress Score", f"{pol}% {pol_emoji}")
            with col3:
                env = job_to_show.get('sensory_overload_index')
                env_emoji = "🔴" if env and env >= 50 else ("🟢" if env else "")
                st.metric("Environmental Stress Index", f"{env}% {env_emoji}")

            st.markdown("#### **Evaluation Axes Breakdown**")
            st.json({
                "Environmental & Biological Guardrails (30%)": f"{job_to_show.get('score_environment_guardrails') or 0} pts",
                "Technical & Creative Autonomy (25%)": f"{job_to_show.get('score_technical_autonomy') or 0} pts",
                "Domain Relevance (20%)": f"{job_to_show.get('score_domain_relevance') or 0} pts",
                "Compensation & Capital Potential (15%)": f"{job_to_show.get('score_compensation_potential') or 0} pts",
                "Future-Proofing & Domain Growth (10%)": f"{job_to_show.get('score_future_mobility') or 0} pts"
            })

            st.markdown("#### **Workplace & Strategic Assessment**")
            st.info(f"**Workplace Stress & Politics Risk:**\n{job_to_show.get('biological_stress_risk') or 'N/A'}")
            st.success(f"**Strategic Career Value:**\n{job_to_show.get('strategic_value') or 'N/A'}")
            
            st.markdown(f"**Target Application Strategy:**")
            st.write(f"📝 **Recommended CV:** `{job_to_show.get('recommended_cv_version') or 'N/A'}`")
            st.write(f"🚀 **Next Action:** `{job_to_show.get('next_action') or 'N/A'}`")
        else:
            st.info("No evaluated jobs in view. The daily automation pipeline automatically processes imported unread email jobs.")

with tab_add_job:
    st.subheader("➕ Import a New Job Advertisement")
    st.write("Add raw jobs manually to the Postgres Vault. The daily evaluation cron job will automatically score and route them.")
    if "manual_import_success" in st.session_state:
        st.success(st.session_state["manual_import_success"])
        del st.session_state["manual_import_success"]

    with st.form("custom_job_form"):
        title = st.text_input("Job Title", "Principal AI Architect")
        company = st.text_input("Company Name", "Novartis Pharmaceuticals")
        source = st.selectbox("Source Portal", ["LinkedIn", "MyCareersFuture", "eFinancialCareers", "Gmail"])
        salary = st.text_input("Salary Range Indicator", "SGD 22,000 - SGD 26,000 / month")
        location = st.text_input("Location", "Singapore (Remote)")
        careers_url = st.text_input("Careers Portal Direct Link (Verification)", "https://www.novartis.com/careers")
        desc = st.text_area("Job Description Raw text", "Paste raw details here...")
        
        submitted = st.form_submit_button("Import & Save to Postgres Vault")
        if submitted:
            new_job = {
                "title": title,
                "company": company,
                "source": source,
                "salaryRange": salary,
                "postedDate": str(datetime.date.today()),
                "location": location,
                "careers_portal_url": careers_url if careers_url else f"https://www.{company.lower().replace(' ', '')}.com/careers",
                "description": desc
            }
            if save_new_job_to_db(new_job):
                st.session_state["manual_import_success"] = f"✅ Successfully added '{title}' to the Staging Vault! It is now pending evaluation. To evaluate it immediately, click the '🧠 2. Run LLM Evaluation & Processing' button in the sidebar."
                st.rerun()

with tab_linkedin:
    st.subheader("🔗 Import Saved LinkedIn Jobs")
    st.write("Sync your saved LinkedIn jobs and stage them in your Postgres database for evaluation. Auto-unsave is disabled to keep ingestion read-only.")
    
    col_auto, col_manual = st.columns(2)
    
    with col_auto:
        st.markdown("### 🤖 Option A: Headless Auto-Sync")
        st.warning("Auto-sync/unsave is disabled by policy to keep source ingestion non-destructive.")
        st.caption("Use Option B (manual export and upload) for read-only ingestion.")

    with col_manual:
        st.markdown("### 📋 Option B: Manual Export & Upload")
        st.write("If your cookie expires or you prefer manual control, run the browser console script below and upload the exported JSON file.")
        script_code = r"""(async function extractSavedJobs() {
  if (window._jobScraperRunning) {
    alert("Job scraper is already running! Please wait for it to finish or refresh the page.");
    return;
  }
  window._jobScraperRunning = true;

  // Create visual overlay for progress
  const overlay = document.createElement('div');
  overlay.style.cssText = 'position:fixed;top:0;left:0;width:100%;height:100%;background:rgba(0,0,0,0.8);color:white;z-index:999999;display:flex;flex-direction:column;align-items:center;justify-content:center;font-family:sans-serif;font-size:18px;';
  const statusText = document.createElement('div');
  statusText.innerText = "🚀 Starting LinkedIn Saved Jobs extraction...";
  const progressText = document.createElement('div');
  progressText.style.marginTop = '10px';
  progressText.style.fontSize = '24px';
  progressText.style.fontWeight = 'bold';
  overlay.appendChild(statusText);
  overlay.appendChild(progressText);
  document.body.appendChild(overlay);

  try {
    const uniqueJobs = [];
    const processedUrls = new Set();
    let pageNum = 1;
    let hasNextPage = true;
    
    while (hasNextPage && pageNum <= 40) {
      statusText.innerText = `📄 Scanning Page ${pageNum}...`;
      
      // Scroll to bottom to ensure elements render
      window.scrollTo(0, document.body.scrollHeight);
      await new Promise(r => setTimeout(r, 2000));
      
      // Find all job links on this page
      const jobLinks = Array.from(document.querySelectorAll('a[href*="/jobs/view/"]'));
      let pageCount = 0;
      
      for (const a of jobLinks) {
        const title = (a.innerText || "").trim();
        if (!title || title.length < 3) continue;
        
        let jobId = '';
        const match = a.href.match(/\/jobs\/view\/(\d+)/);
        if (match) jobId = match[1];
        
        if (jobId) {
          const standardUrl = `https://www.linkedin.com/jobs/view/${jobId}/`;
          
          if (!processedUrls.has(standardUrl)) {
            processedUrls.add(standardUrl);
            
            const container = a.closest('li') || a.closest('.entity-list-item') || a.closest('div');
            let company = 'Unknown Company';
            let location = 'Singapore';
            
            if (container) {
              const companyEl = container.querySelector('.entity-list-item__subtitle, .reusable-search__result-subtitle, .job-card-container__company-name');
              if (companyEl) company = companyEl.innerText.trim();
              
              const locationEl = container.querySelector('.entity-list-item__caption, .reusable-search__result-caption, .job-card-container__metadata-item');
              if (locationEl) location = locationEl.innerText.trim();
              
              if (company === 'Unknown Company') {
                const innerSpans = Array.from(container.querySelectorAll('span, div, p'));
                for (const span of innerSpans) {
                  const t = span.innerText.trim();
                  if (t.includes('·') && !t.includes('\n')) {
                    const parts = t.split('·');
                    company = parts[0].trim();
                    location = parts[1].trim();
                    break;
                  }
                }
              }
            }
            
            uniqueJobs.push({ title, company, url: standardUrl, location });
            pageCount++;
          }
        }
      }
      
      progressText.innerText = `Found ${uniqueJobs.length} unique jobs so far.`;
      
      // Find Next button
      const nextBtn = document.querySelector('.artdeco-pagination__button--next') || 
                      Array.from(document.querySelectorAll('button, a')).find(el => {
                        const text = el.innerText.trim().toLowerCase();
                        return (text.includes('next') || el.ariaLabel?.toLowerCase().includes('next')) && !el.disabled && !el.classList.contains('disabled');
                      });
                      
      if (nextBtn && !nextBtn.disabled && !nextBtn.classList.contains('artdeco-button--disabled')) {
        statusText.innerText = `➡️ Moving to Next page...`;
        nextBtn.click();
        pageNum++;
        await new Promise(r => setTimeout(r, 3000));
      } else {
        statusText.innerText = `🏁 No more pages found. Ending scan.`;
        hasNextPage = false;
      }
    }
    
    if (uniqueJobs.length === 0) {
      alert("⚠️ No saved jobs identified. Make sure you are on the 'Saved' tab of your Job Tracker.");
      window._jobScraperRunning = false;
      document.body.removeChild(overlay);
      return;
    }
    
    statusText.innerText = "🧠 Fetching job descriptions...";
    const finalizedJobs = [];
    
    // Robust XHR wrapper to bypass broken extensions intercepting fetch
    function robustGet(url) {
      return new Promise((resolve, reject) => {
        const xhr = new XMLHttpRequest();
        xhr.open('GET', url);
        xhr.setRequestHeader('Accept', 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8');
        xhr.withCredentials = true;
        xhr.onload = () => resolve(xhr.responseText);
        xhr.onerror = () => reject(new Error('XHR Error'));
        xhr.send();
      });
    }
    
    // Sequential fetching to avoid rate limits
    for (let i = 0; i < uniqueJobs.length; i++) {
      const job = uniqueJobs[i];
      progressText.innerText = `Fetching description ${i + 1} of ${uniqueJobs.length}...\n${job.title}`;
      
      try {
        const html = await robustGet(job.url);
        const parser = new DOMParser();
        const doc = parser.parseFromString(html, 'text/html');
        const descEl = doc.querySelector('.jobs-description-content') || 
                      doc.querySelector('.show-more-less-html__markup') || 
                      doc.querySelector('[id^="job-details"]') || 
                      doc.querySelector('.jobs-box__html-content') ||
                      doc.querySelector('.jobs-description');
        const description = descEl ? descEl.innerText.trim() : '';
        
        finalizedJobs.push({
          title: job.title,
          company: job.company,
          url: job.url,
          location: job.location,
          description: description || "Full description not available. Please visit job link to apply."
        });
      } catch (err) {
        console.error(`❌ Failed: ${job.title}`, err);
        finalizedJobs.push({
          title: job.title,
          company: job.company,
          url: job.url,
          location: job.location,
          description: "Failed to fetch description automatically."
        });
      }
      
      // Delay to avoid LinkedIn 429 rate limit or SDUI oops page
      await new Promise(r => setTimeout(r, 2000));
    }
    
    statusText.innerText = "🎉 All done! Downloading JSON...";
    progressText.innerText = "";
    
    const dataStr = "data:text/json;charset=utf-8," + encodeURIComponent(JSON.stringify(finalizedJobs, null, 2));
    const downloadAnchor = document.createElement('a');
    downloadAnchor.setAttribute("href", dataStr);
    downloadAnchor.setAttribute("download", "linkedin_saved_jobs.json");
    document.body.appendChild(downloadAnchor);
    downloadAnchor.click();
    downloadAnchor.remove();
    
    await new Promise(r => setTimeout(r, 2000));
  } catch (err) {
    console.error("Critical error during extraction:", err);
    alert("Extraction failed. See console for details.");
  } finally {
    window._jobScraperRunning = false;
    document.body.removeChild(overlay);
  }
})();"""
        st.code(script_code, language="javascript")
        
        uploaded_file = st.file_uploader("Upload your exported 'linkedin_saved_jobs.json' file", type=["json"])
        
        if uploaded_file is not None:
            try:
                jobs_data = json.load(uploaded_file)
                if not isinstance(jobs_data, list):
                    st.error("Invalid file structure. Root must be a JSON array.")
                else:
                    st.success(f"Successfully parsed JSON. Found **{len(jobs_data)}** jobs in file.")
                    
                    # Show preview
                    st.markdown("#### **Jobs Preview**")
                    preview_df = pd.DataFrame([
                        {"Title": j.get("title"), "Company": j.get("company"), "Location": j.get("location"), "URL": j.get("url")}
                        for j in jobs_data[:10]
                    ])
                    st.dataframe(preview_df, use_container_width=True)
                    if len(jobs_data) > 10:
                        st.write(f"*... and {len(jobs_data) - 10} more jobs.*")
                    
                    # Ingest button
                    if st.button("📥 Import All Saved Jobs to Staging Vault"):
                        with st.spinner("Inserting jobs into staging table..."):
                            inserted, skipped = ingest_linkedin_saved_json(jobs_data)
                            st.success(f"✅ Ingestion complete! **{inserted}** new jobs imported, **{skipped}** skipped as duplicates/invalid.")
                            st.rerun()
            except Exception as e:
                st.error(f"Error reading JSON file: {e}")

with tab_analytics:
    st.subheader("🔥 Company Autonomy & Culture Analytics")
    st.write("Consolidated employer ratings compiled directly from Postgres database tables.")

    analytics_data = fetch_company_analytics_from_db()

    if analytics_data:
        df = pd.DataFrame(analytics_data)
        
        # Map safety verdicts
        def get_verdict(row):
            st = row.get("Status")
            if row.get("Approved"):
                return "🟢 HIGH AUTONOMY / OPTIMAL FOCUS"
            if st in ("STRONG MATCH", "PRIORITY_APPLY", "HIGH_FIT_HIGH_RISK"):
                return "✅ PRIORITY"
            elif st in ("REVIEW REQUIRED", "APPLY_AFTER_VERIFICATION"):
                return "⚠️ REVIEW REQUIRED"
            elif row.get("Toxic"):
                return "🔴 HIGH POLITICS / BUREAUCRATIC"
            else:
                return "🟡 REVIEW REQUIRED"
                
        df["Safety Verdict"] = df.apply(get_verdict, axis=1)
        
        # Display table
        st.dataframe(df[["Company", "Industry", "Avg Autonomy Score", "Avg Politics Score", "Avg Sensory Index", "Avg Focus Score", "Safety Verdict"]], use_container_width=True)
        
        st.markdown("### 🏆 Top High-Autonomy Environments vs. ⚠️ Stress Alerts")
        st.bar_chart(df.set_index("Company")[["Avg Autonomy Score", "Avg Politics Score"]])
    else:
        st.info("No compiled analytics are available. Please run the evaluation engine pipeline to score listings.")

with tab_cv:
    st.subheader("📄 Canonical Documents")
    st.write("Generate CV and cover letter from canonical job/version records and the single master_profile evidence ledger.")

    profile_source = None
    profile_data = None
    if os.environ.get("MASTER_PROFILE_JSON"):
        try:
            profile_data = json.loads(os.environ["MASTER_PROFILE_JSON"])
            profile_source = "MASTER_PROFILE_JSON"
        except Exception as e:
            st.error(f"Invalid MASTER_PROFILE_JSON: {e}")
    elif os.path.exists("master_profile.json"):
        try:
            with open("master_profile.json", "r", encoding="utf-8") as f:
                profile_data = json.load(f)
            profile_source = "master_profile.json"
        except Exception as e:
            st.error(f"Failed to parse master_profile.json: {e}")
    else:
        st.error("master_profile evidence ledger is missing. Provide MASTER_PROFILE_JSON or master_profile.json.")

    if profile_data:
        facts = profile_data.get("profile_facts") or profile_data.get("facts") or []
        fact_ids = [f.get("id") for f in facts if isinstance(f, dict) and f.get("id")]
        st.caption(f"Evidence source: {profile_source} | facts loaded: {len(fact_ids)}")
        if len(fact_ids) == 0:
            st.error("No profile fact IDs found in the master profile ledger. Document generation is disabled.")

    eligible_jobs = [j for j in jobs_list if j.get("id") and j.get("job_version_id") and j.get("status") in ("AI_EVALUATED", "QUEUED_FOR_AI", "LANE_ROUTED", "PREQUALIFIED", "NEEDS_VERIFICATION")]

    if not eligible_jobs:
        st.info("No canonical shortlist jobs with version IDs are currently available.")
    elif not profile_data or len((profile_data.get("profile_facts") or profile_data.get("facts") or [])) == 0:
        st.info("Fix the master profile ledger first, then generate documents.")
    else:
        def _job_label(j):
            return f"{j.get('company')} - {j.get('title')} [{j.get('status')}] ({str(j.get('job_version_id'))[:8]})"

        options = {_job_label(j): j for j in eligible_jobs}
        selected_label = st.selectbox("Select canonical job/version", list(options.keys()))
        selected_job = options[selected_label]

        st.markdown(f"**Canonical Job ID:** {selected_job.get('id')}")
        st.markdown(f"**Job Version ID:** {selected_job.get('job_version_id')}")
        st.markdown(f"**Company:** {selected_job.get('company')} | **Role:** {selected_job.get('title')}")
        if selected_job.get("careers_portal_url"):
            st.markdown(f"🔗 [View Posting]({selected_job.get('careers_portal_url')})")

        with st.expander("Preview job description"):
            description = selected_job.get("description") or ""
            if isinstance(description, dict):
                st.json(description)
            else:
                st.text_area("Description", str(description), height=220, disabled=True)

        btn_col1, btn_col2 = st.columns(2)
        with btn_col1:
            if st.button("Generate customized CV", use_container_width=True):
                with st.spinner("Generating CV from canonical job/version..."):
                    try:
                        run_checked_command(
                            [
                                "npx",
                                "tsx",
                                "scripts/generate_cv.ts",
                                str(selected_job.get("id")),
                                str(selected_job.get("job_version_id"))
                            ],
                            "generate_cv"
                        )
                        st.success("CV generation completed. See scripts/exports for artifacts.")
                    except subprocess.CalledProcessError as e:
                        st.error(f"CV generation failed with exit code {e.returncode}.")
                        if e.stdout:
                            st.code(e.stdout, language="text")
                        if e.stderr:
                            st.error(e.stderr)

        with btn_col2:
            if st.button("Generate cover letter", use_container_width=True):
                with st.spinner("Generating cover letter from canonical job/version..."):
                    try:
                        run_checked_command(
                            [
                                "npx",
                                "tsx",
                                "scripts/generate_cover_letter.ts",
                                str(selected_job.get("id")),
                                str(selected_job.get("job_version_id"))
                            ],
                            "generate_cover_letter"
                        )
                        st.success("Cover letter generation completed. See scripts/exports for artifacts.")
                    except subprocess.CalledProcessError as e:
                        st.error(f"Cover letter generation failed with exit code {e.returncode}.")
                        if e.stdout:
                            st.code(e.stdout, language="text")
                        if e.stderr:
                            st.error(e.stderr)

# Footer section
st.markdown("---")
st.markdown("<p class='disclaimer'>Job Decision Engine v4.0 • Powered by Neon Postgres & GitHub Actions Automation</p>", unsafe_allow_html=True)
