import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# Selected design system: standard_business_brief.
NAVY = RGBColor(21, 58, 91)
BLUE = RGBColor(31, 90, 122)
INK = RGBColor(35, 42, 48)
MUTED = RGBColor(92, 101, 110)
LIGHT = RGBColor(225, 231, 236)
WHITE = RGBColor(255, 255, 255)
FONT = "Calibri"

def set_font(run, size=None, bold=None, italic=None, color=INK, name=FONT):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color

def set_cell_or_para_shading(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_paragraph_border_bottom(paragraph, color="D7E0E8", size="8", space="2"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)

def add_page_number(paragraph, name):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run(f"{name}  |  ")
    set_font(r, 8.5, color=MUTED)
    r = paragraph.add_run("Page ")
    set_font(r, 8.5, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()
    r._r.extend([fld_begin, instr, fld_sep, fld_end])

def create_bullet_num_id(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 1
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "260")
    ppr.extend([tabs, ind])
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    rpr.append(rfonts)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, ppr, rpr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id

def add_bullet(doc, num_id, text, bold_lead=None, after=2.4):
    p = doc.add_paragraph(style="Resume Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.03
    p.paragraph_format.keep_together = True
    num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, numid])
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_font(r, 9.6, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_font(r, 9.6)
    else:
        r = p.add_run(text)
        set_font(r, 9.6)
    return p

def add_section(doc, title, before=8, after=4):
    p = doc.add_paragraph(style="Resume Section")
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title.upper())
    set_font(r, 10.8, bold=True, color=NAVY)
    set_paragraph_border_bottom(p, color="D6E1EA", size="5", space="2")
    return p

def add_role(doc, company, period, title, scope=None, location=None, before=5):
    p = doc.add_paragraph(style="Resume Role")
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(6.5), alignment=2)
    r = p.add_run(company)
    set_font(r, 10.1, bold=True, color=INK)
    r = p.add_run("\t" + period)
    set_font(r, 9.3, bold=True, color=MUTED)

    p2 = doc.add_paragraph(style="Resume Role Title")
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.keep_with_next = True
    r = p2.add_run(title)
    set_font(r, 9.6, bold=True, color=BLUE)
    if location:
        r = p2.add_run(f"  |  {location}")
        set_font(r, 9.2, color=MUTED)
    if scope and str(scope).strip().lower() not in ["null", "none", ""]:
        p3 = doc.add_paragraph(style="Resume Scope")
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after = Pt(2)
        p3.paragraph_format.keep_with_next = True
        r = p3.add_run("Functional scope: ")
        set_font(r, 9.2, bold=True, color=MUTED)
        r = p3.add_run(scope)
        set_font(r, 9.2, italic=True, color=MUTED)

def add_label_para(doc, label, text, after=2.5):
    p = doc.add_paragraph(style="Resume Body")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.03
    p.paragraph_format.keep_together = True
    r = p.add_run(label + ": ")
    set_font(r, 9.6, bold=True, color=BLUE)
    r = p.add_run(text)
    set_font(r, 9.6)
    return p

def add_compact_line(doc, bold_text, text, after=1.8):
    p = doc.add_paragraph(style="Resume Body")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.02
    p.paragraph_format.keep_together = True
    r = p.add_run(bold_text)
    set_font(r, 9.4, bold=True)
    r = p.add_run(text)
    set_font(r, 9.4)
    return p

def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    p.paragraph_format.space_after = Pt(0)

def render(json_path, out_path):
    with open(json_path, 'r', encoding='utf-8') as f:
        cv = json.load(f)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    for name in ["Resume Body", "Resume Bullet", "Resume Section", "Resume Role", "Resume Role Title", "Resume Scope"]:
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(10)
        style.font.color.rgb = INK

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    add_page_number(fp, cv.get('candidate', {}).get('name', 'Candidate'))

    bullet_num_id = create_bullet_num_id(doc)

    c = cv.get('candidate', {})
    t = cv.get('target', {})

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(c.get('name', '').upper())
    set_font(r, 24, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(t.get('headline', '').upper())
    set_font(r, 11.2, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(f"{c.get('location', '')}  |  {c.get('phone', '')}  |  {c.get('email', '')}  |  {c.get('linkedin', '')}")
    set_font(r, 9.2, color=INK)
    set_paragraph_border_bottom(p, color="AFC3D3", size="8", space="4")

    add_section(doc, "Executive Profile", before=4, after=4)
    p = doc.add_paragraph(style="Resume Body")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(cv.get('executiveProfile', ''))
    set_font(r, 9.8)

    add_section(doc, f"{t.get('company', 'Target')} Role Alignment", before=5, after=3)
    for ra in cv.get('roleAlignment', []):
        add_label_para(doc, ra.get('label', ''), ra.get('summary', ''))

    add_section(doc, "Professional Experience", before=5, after=3)
    for exp in cv.get('experience', []):
        period = f"{exp.get('startDate', '')} - {exp.get('endDate', '')}"
        add_role(doc, exp.get('company', ''), period, exp.get('formalTitle', ''), scope=exp.get('functionalScope'), location=exp.get('location'), before=3)
        for bullet in exp.get('bullets', []):
            add_bullet(doc, bullet_num_id, bullet.get('text', ''))

    add_section(doc, "Education & Credentials", before=6, after=3)
    for edu in cv.get('education', []):
        add_compact_line(doc, "", edu)
    for cert in cv.get('certifications', []):
        add_compact_line(doc, "", cert)

    add_section(doc, "Technical & Domain Skills", before=6, after=3)
    for skill in cv.get('skills', []):
        add_compact_line(doc, "", skill)

    doc.save(out_path)
    print(f"Saved: {out_path}")

if __name__ == "__main__":
    render(sys.argv[1], sys.argv[2])
